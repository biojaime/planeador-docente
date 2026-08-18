import streamlit as st
import json
import os
import re
import sys
from datetime import timedelta, date, datetime
import io
import base64
from streamlit_quill import st_quill
import streamlit.components.v1 as components
from PIL import Image
import hashlib
import tempfile
import uuid
import sqlite3
# removed requests-based login integration per user request

# --- ReportLab Imports ---
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage, PageBreak
from reportlab.platypus import ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus.doctemplate import Indenter
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Configuration ---
st.set_page_config(page_title="Planeador Docente IMM", page_icon="📝", layout="wide")

# --- Helper Functions ---

def _get_full_path(path):
    """
    Get the absolute path for a file, handling both script execution and PyInstaller (if used later).
    Checks current directory first.
    """
    cwd_path = os.path.join(os.getcwd(), path)
    if os.path.exists(cwd_path):
        return cwd_path
    
    alt_path = os.path.join(r"C:\Users\jaime\Documents\My planer", path)
    if os.path.exists(alt_path):
        return alt_path

    try:
        base = sys._MEIPASS
    except AttributeError:
        base = os.path.dirname(os.path.abspath(__file__))
    
    full = os.path.join(base, path)
    return full

def get_image_base64(path):
    """Encodes an image to base64 for embedding in HTML."""
    full_path = _get_full_path(path)
    if os.path.exists(full_path):
        with open(full_path, "rb") as f:
            data = f.read()
        return base64.b64encode(data).decode()
    return ""

def html_to_reportlab(html_text):
    """
    Convert Quill HTML to ReportLab XML tags for use inside a Paragraph.
    Handles <b>, <i>, <u>, <strong>, <em>, <span style=bold/italic>, <br/>, <p>.
    List items (<li>) are rendered as bullet prefix lines.
    """
    if not html_text:
        return ""

    text = str(html_text)

    # Normalize self-closing br variants
    text = text.replace('<br>', '<br/>').replace('<br />', '<br/>')

    # Paragraph tags: opening <p ...> stripped, closing </p> becomes double break
    text = re.sub(r'<p[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'</p>', '<br/>', text, flags=re.IGNORECASE)

    # Bold: <strong> and <span style="font-weight: bold">
    text = re.sub(
        r'<span[^>]*style="[^"]*font-weight\s*:\s*bold[^"]*"[^>]*>(.*?)</span>',
        r'<b>\1</b>', text, flags=re.IGNORECASE | re.DOTALL
    )
    text = re.sub(r'<strong>(.*?)</strong>', r'<b>\1</b>', text, flags=re.IGNORECASE | re.DOTALL)

    # Italic: <em> and <span style="font-style: italic">
    text = re.sub(
        r'<span[^>]*style="[^"]*font-style\s*:\s*italic[^"]*"[^>]*>(.*?)</span>',
        r'<i>\1</i>', text, flags=re.IGNORECASE | re.DOTALL
    )
    text = re.sub(r'<em>(.*?)</em>', r'<i>\1</i>', text, flags=re.IGNORECASE | re.DOTALL)

    # Strip remaining <span> tags (keep content)
    text = re.sub(r'<span[^>]*>', '', text, flags=re.IGNORECASE)
    text = text.replace('</span>', '')

    # List items: convert to bullet prefix (for Paragraph context)
    text = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1<br/>', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'</?[uo]l[^>]*>', '', text, flags=re.IGNORECASE)

    # Strip any remaining unknown tags while preserving ReportLab-safe ones
    # Keep <b>, </b>, <i>, </i>, <u>, </u>, <br/>
    text = re.sub(r'<(?!/?(?:b|i|u|br/))[^>]+>', '', text, flags=re.IGNORECASE)

    # Clean up leading/trailing breaks
    text = text.strip()
    while text.startswith('<br/>'):
        text = text[5:].strip()
    while text.endswith('<br/>'):
        text = text[:-5].strip()

    return text


def html_to_flowables(html_text, styles, bullet_indent=18):
    """
    Convert Quill-like HTML into a list of ReportLab flowables.
    Handles: <p>, <br/>, <b>, <i>, <u>, <strong>, <em>, <ul>, <ol>, <li>, <span style>.
    Returns a list of Paragraph and ListFlowable objects preserving formatting.
    """
    if not html_text:
        return []

    s = str(html_text).replace('\r', '')
    # Normalize br variants
    s = s.replace('<br>', '<br/>').replace('<br />', '<br/>')

    flowables = []
    # Split on list boundaries so we can handle them separately
    tokens = re.split(r'(<ul[^>]*>|</ul>|<ol[^>]*>|</ol>)', s, flags=re.IGNORECASE)

    body_style = styles.get('Normal')
    para_style = ParagraphStyle(
        'FlowBody',
        parent=body_style,
        spaceAfter=4,
        leading=14,
    )
    list_item_style = ParagraphStyle(
        'FlowListItem',
        parent=body_style,
        leading=14,
    )

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        low = tok.lower().strip() if tok else ''

        if re.match(r'<ul[^>]*>', low, re.IGNORECASE):
            # Unordered list block — collect until </ul>
            i += 1
            inner = ''
            while i < len(tokens) and not re.match(r'</ul>', tokens[i].strip(), re.IGNORECASE):
                inner += tokens[i]
                i += 1
            li_items = re.findall(r'<li[^>]*>(.*?)</li>', inner, flags=re.IGNORECASE | re.DOTALL)
            list_items = []
            for li in li_items:
                txt = html_to_reportlab(li)
                if txt:
                    list_items.append(ListItem(Paragraph(txt, list_item_style), leftIndent=bullet_indent))
            if list_items:
                flowables.append(ListFlowable(list_items, bulletType='bullet', leftIndent=bullet_indent))
            i += 1  # skip </ul>

        elif re.match(r'<ol[^>]*>', low, re.IGNORECASE):
            # Ordered list block — collect until </ol>
            i += 1
            inner = ''
            while i < len(tokens) and not re.match(r'</ol>', tokens[i].strip(), re.IGNORECASE):
                inner += tokens[i]
                i += 1
            li_items = re.findall(r'<li[^>]*>(.*?)</li>', inner, flags=re.IGNORECASE | re.DOTALL)
            list_items = []
            for li in li_items:
                txt = html_to_reportlab(li)
                if txt:
                    list_items.append(ListItem(Paragraph(txt, list_item_style), leftIndent=bullet_indent))
            if list_items:
                flowables.append(ListFlowable(list_items, bulletType='1', leftIndent=bullet_indent))
            i += 1  # skip </ol>

        else:
            # Regular HTML block — may contain <p> and <br/>
            if tok and tok.strip():
                # Split on <p> tags first to get paragraph blocks
                # Replace </p> with sentinel, then split
                chunk = tok
                chunk = re.sub(r'<p[^>]*>', '\x00P_OPEN\x00', chunk, flags=re.IGNORECASE)
                chunk = re.sub(r'</p>', '\x00P_CLOSE\x00', chunk, flags=re.IGNORECASE)
                segments = chunk.split('\x00P_OPEN\x00')
                for seg in segments:
                    # Each seg may end with \x00P_CLOSE\x00
                    seg = seg.replace('\x00P_CLOSE\x00', '').strip()
                    if not seg:
                        continue
                    # Now split on double (or more) <br/> to get sub-paragraphs
                    sub_parts = re.split(r'(?:<br\s*/>\s*){2,}', seg)
                    for part in sub_parts:
                        # Single <br/> becomes newline within paragraph
                        part = part.strip()
                        if not part:
                            continue
                        txt = html_to_reportlab(part)
                        if txt:
                            flowables.append(Paragraph(txt, para_style))
            i += 1

    return flowables


def safe_filename(s):
    # Remove or replace characters invalid in Windows filenames
    s = re.sub(r'[\\/:*?"<>|]', '-', str(s))
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def _embed_image_to_pdf(path, content_list, page_width):
    if path and os.path.exists(path):
        try:
            flat_path = flatten_image(path)
            img = RLImage(flat_path, width=page_width, height=4*inch, kind='proportional')
            content_list.append(img)
            content_list.append(Spacer(1, 0.1 * inch))
        except Exception as e:
            content_list.append(Paragraph(f"<i>[Error al cargar imagen: {os.path.basename(path)}]</i>", getSampleStyleSheet()['Italic']))


def cleanup_temp_rubric_files():
    active_paths = set()
    active_paths.update(st.session_state.get('abpj_rubrica_paths', []))
    for day in st.session_state.get('daily_plan_data', {}).values():
        active_paths.update(day.get('rubrica_paths', []))

    for filename in os.listdir(os.getcwd()):
        if filename.startswith('temp_rubric_'):
            full_path = os.path.abspath(filename)
            if full_path not in active_paths:
                try:
                    os.remove(full_path)
                except Exception:
                    pass


def html_to_plain_text(html_text):
    if not html_text:
        return ''
    text = str(html_text)
    text = text.replace('<br>', '\n').replace('<br/>', '\n').replace('<br />', '\n')
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def fill_docx_cell_from_html(cell, html_text):
    content = str(html_text or '')
    if not content.strip():
        return

    content = content.replace('<br>', '<br/>').replace('<br />', '<br/>')
    tokens = re.split(r'(<\/?[^>]+>)', content)

    paragraph = cell.paragraphs[0]
    paragraph.text = ''
    bold = italic = underline = False
    list_style = None

    for token in tokens:
        if not token:
            continue
        if token.lower() == '<p>':
            if paragraph.text or paragraph.runs:
                paragraph = cell.add_paragraph()
            continue
        if token.lower() == '</p>':
            paragraph = cell.add_paragraph(style=list_style)
            continue
        if token.lower() == '<ul>':
            list_style = 'List Bullet'
            continue
        if token.lower() == '</ul>':
            list_style = None
            paragraph = cell.add_paragraph()
            continue
        if token.lower() == '<ol>':
            list_style = 'List Number'
            continue
        if token.lower() == '</ol>':
            list_style = None
            paragraph = cell.add_paragraph()
            continue
        if token.lower() == '<li>':
            paragraph = cell.add_paragraph(style=list_style or 'List Bullet')
            continue
        if token.lower() == '</li>':
            paragraph = cell.add_paragraph(style=list_style)
            continue
        if token.lower() in ('<strong>', '<b>'):
            bold = True
            continue
        if token.lower() in ('</strong>', '</b>'):
            bold = False
            continue
        if token.lower() in ('<em>', '<i>'):
            italic = True
            continue
        if token.lower() in ('</em>', '</i>'):
            italic = False
            continue
        if token.lower() == '<u>':
            underline = True
            continue
        if token.lower() == '</u>':
            underline = False
            continue
        if token.lower() == '<br/>' or token.lower() == '<br />':
            paragraph.add_run().add_break()
            continue

        run = paragraph.add_run(token)
        run.bold = bold
        run.italic = italic
        run.underline = underline


def set_docx_landscape(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def flatten_image(src_path, force_recreate=False):
    """If image has alpha/transparency, flatten it on white background and return a path to a non-alpha image.
    Caches the flattened image next to the original with suffix `_flat.jpg` to avoid repeating work.
    """
    try:
        dirname = os.path.dirname(src_path)
        base = os.path.splitext(os.path.basename(src_path))[0]
        out_name = os.path.join(dirname, f"{base}_flat.jpg")

        if os.path.exists(out_name) and not force_recreate:
            if os.path.getmtime(out_name) >= os.path.getmtime(src_path):
                return out_name

        with Image.open(src_path) as im:
            if im.mode in ("RGBA", "LA") or (im.mode == "P" and "transparency" in im.info):
                bg = Image.new("RGB", im.size, (255, 255, 255))
                if im.mode != "RGBA":
                    im = im.convert("RGBA")
                bg.paste(im, mask=im.split()[3])
                bg.save(out_name, "JPEG", quality=95)
                return out_name
            else:
                im.convert("RGB").save(out_name, "JPEG", quality=95)
                return out_name
    except Exception:
        return src_path

def parse_date(date_str):
    """Parses a date string trying ISO format first, then DD/MM/YYYY."""
    if not date_str:
        return date.today()
    try:
        return date.fromisoformat(date_str)
    except ValueError:
        try:
            return datetime.strptime(date_str, "%d/%m/%Y").date()
        except ValueError:
            return date.today()

# --- Load PDA Database ---
@st.cache_data
def load_pda_database():
    import csv
    csv_path = _get_full_path("PProceso de Desarrollo de Aprendizaje (PDA).csv")
    if not os.path.exists(csv_path):
        return []
    records = []
    try:
        with open(csv_path, mode="r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            header = next(reader)
            col_map = {col.strip(): idx for idx, col in enumerate(header)}
            
            idx_campo = col_map.get("Campo Formativo", 0)
            idx_materia = col_map.get("Materia", 1)
            idx_grado = col_map.get("Grado", 2)
            idx_contenido = col_map.get("Contenido", 3)
            idx_pda = col_map.get("Proceso de Desarrollo de Aprendizaje (PDA)", 4)
            
            for row in reader:
                if len(row) > max(idx_campo, idx_materia, idx_grado, idx_contenido, idx_pda):
                    records.append({
                        "campo": row[idx_campo].strip(),
                        "materia": row[idx_materia].strip(),
                        "grado": row[idx_grado].strip(),
                        "contenido": row[idx_contenido].strip(),
                        "pda": row[idx_pda].strip()
                    })
    except Exception as e:
        st.error(f"Error al cargar la base de datos de PDA: {e}")
    return records

def map_subject_and_grade(materia, curso_grado):
    grado_map = {"1ro": "1\u00ba", "2do": "2\u00ba", "3ro": "3\u00ba"}
    target_grado = grado_map.get(curso_grado, "1\u00ba")
    
    materia_clean = materia.strip()
    
    if materia_clean.endswith(" III"):
        base_subject = materia_clean[:-4].strip()
        target_grado = "3\u00ba"
    elif materia_clean.endswith(" II"):
        base_subject = materia_clean[:-3].strip()
        target_grado = "2\u00ba"
    elif materia_clean.endswith(" I"):
        base_subject = materia_clean[:-2].strip()
        target_grado = "1\u00ba"
    else:
        base_subject = materia_clean
        
    materia_translation = {
        "Matematicas": "Matem\u00e1ticas",
        "Español": "Espa\u00f1ol",
        "Educación Civica y Etica": "Formaci\u00f3n C\u00edvica y \u00c9tica",
        "Ingles": "Ingl\u00e9s",
        "Informatica": "Tecnolog\u00eda",
        "Historia": "Historia",
        "Educación Fisica": "Educaci\u00f3n F\u00edsica",
        "Biología": "Biolog\u00eda",
        "Fisica": "F\u00edsica",
        "Quimica": "Qu\u00edmica",
    }
    
    if base_subject == "Ciencias":
        if target_grado == "1\u00ba":
            return "Biolog\u00eda", "1\u00ba"
        elif target_grado == "2\u00ba":
            return "F\u00edsica", "2\u00ba"
        elif target_grado == "3\u00ba":
            return "Qu\u00edmica", "3\u00ba"
            
    if base_subject == "Biología":
        return "Biolog\u00eda", "1\u00ba"
    if base_subject == "Fisica":
        return "F\u00edsica", "2\u00ba"
    if base_subject == "Quimica":
        return "Qu\u00edmica", "3\u00ba"
        
    csv_subject = materia_translation.get(base_subject, base_subject)
    return csv_subject, target_grado

def get_pdas_for_selection(materia, curso_grado):
    records = load_pda_database()
    csv_subject, target_grado = map_subject_and_grade(materia, curso_grado)
    
    filtered = []
    for r in records:
        def normalize(text):
            import unicodedata
            return "".join(c for c in unicodedata.normalize('NFD', text) if unicodedata.category(c) != 'Mn').lower().strip()
            
        if normalize(r["materia"]) == normalize(csv_subject) and normalize(r["grado"]) == normalize(target_grado):
            filtered.append(r)
    return filtered

def safe_rerun():
    """Trigger Streamlit rerun safely across versions."""
    try:
        st.rerun()
    except Exception:
        try:
            rerun_fn = getattr(st, "experimental_rerun", None)
            if rerun_fn:
                rerun_fn()
        except Exception:
            pass

def process_uploaded_rubrics(uploaded_files, current_paths, key_prefix):
    """Processes uploaded rubric files, preventing duplicate entries across reruns."""
    if "processed_upload_hashes" not in st.session_state or not isinstance(st.session_state.processed_upload_hashes, set):
        st.session_state.processed_upload_hashes = set()
    if "path_hash_map" not in st.session_state or not isinstance(st.session_state.path_hash_map, dict):
        st.session_state.path_hash_map = {}

    if not uploaded_files:
        return list(current_paths or [])

    if not isinstance(uploaded_files, list):
        uploaded_files = [uploaded_files]

    updated_paths = list(current_paths or [])

    for idx, uploaded in enumerate(uploaded_files):
        try:
            buf = uploaded.getbuffer()
            file_hash = hashlib.sha256(buf).hexdigest()

            if file_hash in st.session_state.processed_upload_hashes:
                continue

            safe_name = re.sub(r'[^a-zA-Z0-9_.-]', '_', uploaded.name)
            temp_path = os.path.abspath(f"temp_rubric_{key_prefix}_{idx}_{safe_name}")
            with open(temp_path, "wb") as f:
                f.write(buf)

            if temp_path not in updated_paths:
                updated_paths.append(temp_path)

            st.session_state.processed_upload_hashes.add(file_hash)
            st.session_state.path_hash_map[temp_path] = file_hash
        except Exception as e:
            st.error(f"Error al procesar imagen {uploaded.name}: {e}")

    return updated_paths

def remove_rubric_image(image_path, paths_list):
    """Safely removes a rubric image file from disk and session state without crashing."""
    if image_path in paths_list:
        paths_list.remove(image_path)

    path_map = st.session_state.get("path_hash_map", {})
    if image_path in path_map:
        file_hash = path_map.pop(image_path)
        hashes_set = st.session_state.get("processed_upload_hashes")
        if isinstance(hashes_set, set):
            hashes_set.discard(file_hash)

    if os.path.exists(image_path):
        try:
            os.remove(image_path)
        except Exception:
            pass

    safe_rerun()

def get_current_pda_list():
    """Returns a unified list of strings for all active PDAs (both official and custom)."""
    entries = []
    for pda_text in st.session_state.get("pda_official_selected", []):
        if pda_text and str(pda_text).strip():
            entries.append(str(pda_text).strip())

    for custom_text in st.session_state.get("pda_custom_entries", []):
        if custom_text and str(custom_text).strip():
            entries.append(str(custom_text).strip())

    return entries

def get_current_pda():
    """Returns a unified string representation of all active PDAs."""
    entries = get_current_pda_list()
    if not entries:
        return ""
    if len(entries) == 1:
        return entries[0]
    items = []
    for e in entries:
        clean = re.sub(r'<[^>]+>', '', str(e)).strip()
        if clean:
            items.append(f"• {clean}")
    return "<br/>".join(items)

# --- Initialization & State ---

def init_session_state():
    defaults = {
        "logged_in": False,
        "docente_titulo": "Dr.",
        "docente_nombre": "",
        "curso_grado": "1ro",
        "curso_grupos": [],
        "curso_materia": "Matematicas",
        "curso_campo": "Lenguajes",
        "curso_bilingue": False,
        "plan_metodologia": "Seleccione metodología",
        "plan_fecha_inicio": date.today(),
        "plan_fecha_fin": date.today(),
        "plan_dias": [],
        "plan_eje1": "Seleccione eje",
        "plan_eje2": "Seleccione eje",
        "plan_eje3": "Seleccione eje",
        "plan_disc1": "Seleccione materia",
        "plan_disc2": "Seleccione materia",
        "plan_disc3": "Seleccione materia",
        "text_problematica": "",
        "text_pda": "",
        "pda_custom_active": False,
        "pda_selected": "",
        "pda_custom": "",
        "pda_entries": [],
        "pda_official_selected": [],
        "pda_custom_entries": [],
        "processed_upload_hashes": set(),
        "path_hash_map": {},
        "text_objetivos": "",
        "text_perfiles": "",
        "text_producto": "",
        "abpj_presentacion": "",
        "abpj_recoleccion": "",
        "abpj_formulacion": "",
        "abpj_organizacion": "",
        "abpj_experiencia": "",
        "abpj_resultados": "",
        "abpj_materiales": "",
        "abpj_evaluacion": "",
        "abpj_rubrica_paths": [],
        "daily_plan_data": {},
        "last_loaded_file_id": None,
        "quill_key_suffix": 0 # Force Quill refresh
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()


def _compute_data_hash(data_obj):
    try:
        j = json.dumps(data_obj, sort_keys=True, ensure_ascii=False).encode('utf-8')
        return hashlib.sha256(j).hexdigest()
    except Exception:
        return None


# --- Database & Persistence Infrastructure (SQLite) ---

def get_db_path():
    """
    Returns path to SQLite database. Uses environment variable PLANEADOR_DB_PATH if set,
    otherwise defaults to ./data/planeador.db relative to script directory.
    """
    env_path = os.environ.get("PLANEADOR_DB_PATH")
    if env_path:
        return env_path
    base_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_dir, "data", "planeador.db")

def init_db():
    """Initializes SQLite database with WAL journal mode for high concurrency."""
    db_file = get_db_path()
    os.makedirs(os.path.dirname(os.path.abspath(db_file)), exist_ok=True)
    try:
        with sqlite3.connect(db_file, timeout=10) as conn:
            conn.execute("PRAGMA journal_mode=WAL;")
            conn.execute("""
                CREATE TABLE IF NOT EXISTS planeaciones (
                    user_id TEXT PRIMARY KEY,
                    docente_nombre TEXT,
                    data_json TEXT NOT NULL,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                );
            """)
            conn.commit()
    except Exception as e:
        st.error(f"Error al inicializar la base de datos de planeaciones: {e}")

def save_user_planeacion(user_id, data_obj):
    """Saves user planning data to SQLite database."""
    if not user_id:
        return
    init_db()
    json_str = json.dumps(data_obj, ensure_ascii=False, indent=2, default=str)
    docente_name = data_obj.get("docente", {}).get("nombre", "")
    db_file = get_db_path()
    try:
        with sqlite3.connect(db_file, timeout=10) as conn:
            conn.execute("""
                INSERT INTO planeaciones (user_id, docente_nombre, data_json, updated_at)
                VALUES (?, ?, ?, CURRENT_TIMESTAMP)
                ON CONFLICT(user_id) DO UPDATE SET
                    docente_nombre=excluded.docente_nombre,
                    data_json=excluded.data_json,
                    updated_at=CURRENT_TIMESTAMP;
            """, (user_id, docente_name, json_str))
            conn.commit()
    except Exception:
        pass

def load_user_planeacion(user_id):
    """Loads user planning data from SQLite database."""
    if not user_id:
        return None
    init_db()
    db_file = get_db_path()
    try:
        with sqlite3.connect(db_file, timeout=10) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT data_json FROM planeaciones WHERE user_id = ?", (user_id,))
            row = cursor.fetchone()
            if row and row[0]:
                return json.loads(row[0])
    except Exception:
        pass
    return None

def delete_user_planeacion(user_id):
    """Deletes a user's autosave record from SQLite database."""
    if not user_id:
        return
    init_db()
    db_file = get_db_path()
    try:
        with sqlite3.connect(db_file, timeout=10) as conn:
            conn.execute("DELETE FROM planeaciones WHERE user_id = ?", (user_id,))
            conn.commit()
    except Exception:
        pass

def get_user_id():
    """
    Get or create a stable user_id for the current session.
    Checks query params first ('user', 'user_id', 'client_id'), then session state,
    or auto-generates a clean unique key.
    """
    if st.session_state.get("user_id"):
        return st.session_state["user_id"]

    try:
        q_user = st.query_params.get("user") or st.query_params.get("user_id") or st.query_params.get("client_id")
        if q_user and isinstance(q_user, str) and q_user.strip():
            safe_id = re.sub(r'[^a-zA-Z0-9_-]', '', q_user.strip())
            if safe_id:
                st.session_state["user_id"] = safe_id
                return safe_id
    except Exception:
        pass

    new_id = f"user_{uuid.uuid4().hex[:10]}"
    st.session_state["user_id"] = new_id
    try:
        st.query_params["user"] = new_id
    except Exception:
        pass
    return new_id

def autosave_current_data():
    """Save the current planning data to SQLite database when changes are detected."""
    try:
        if not st.session_state.get("autosave_enabled", True):
            return

        data = get_current_data()
        h = _compute_data_hash(data)
        if not h:
            return

        if st.session_state.get("_autosave_hash") != h:
            user_id = get_user_id()
            save_user_planeacion(user_id, data)
            st.session_state["_autosave_hash"] = h
    except Exception:
        pass

# --- Lists ---
LISTA_MATERIAS = ["Matematicas", "Matematicas I", "Matematicas II", "Matematicas III", "Español", "Español I", "Español II", "Español III", "Educación Civica y Etica", "Educación Civica y Etica I", "Educación Civica y Etica II", "Educación Civica y Etica III", "Ingles", "Ingles I", "Ingles II", "Ingles III", "Informatica", "Informatica I", "Informatica II", "Informatica III", "Historia", "Historia I", "Historia II", "Historia III", "Educación Fisica", "Artes", "Ciencias", "Biología", "Fisica", "Quimica"]
LISTA_METODOLOGIA = ["Seleccione metodología", "Aprendizaje Basado en Proyectos (ABPj)", "Aprendizaje Basado en Problemas (ABP)", "STEAM", "Clase invertida (Flipped Classroom)", "Aprendizaje Servicio (ApS)", "Gamificación", "Aprendizaje autodirigido", "Aprendizaje situado", "Aprendizaje entre pares"]
LISTA_EJES = ["Seleccione eje", "Pensamiento Crítico", "Interculturalidad Crítica", "Igualdad de Género", "Vida Saludable", "Apropiación de las Culturas a través de la Lectura y la Escritura", "Artes y Experiencias Estéticas", "Inclusión"]
LISTA_CAMPOS = ["Lenguajes", "Saberes y Pensamiento Científico", "Ética, Naturaleza y Sociedades", "De lo Humano y lo Comunitario"]
LISTA_GRUPOS = ["A", "B", "C", "D", "E", "F"]
LISTA_DIAS = ["Lunes", "Martes", "Miércoles", "Jueves", "Viernes"]

# --- Floating Help Button CSS ---
help_img_b64 = get_image_base64("Help.png")
gemini_img_b64 = get_image_base64("Gemini.png")
deepseek_img_b64 = get_image_base64("DeepSeek.png")

st.markdown(f"""
<style>
.floating-container {{
    position: fixed;
    bottom: 20px;
    right: 20px;
    z-index: 1000;
    display: flex;
    flex-direction: column-reverse;
    align-items: center;
    gap: 10px;
}}
.help-button {{
    width: 60px;
    height: 60px;
    border-radius: 50%;
    background-image: url('data:image/png;base64,{help_img_b64}');
    background-size: cover;
    cursor: pointer;
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    transition: transform 0.3s;
}}
.help-button:hover {{
    transform: scale(1.1);
}}
.popup-icons {{
    display: none;
    flex-direction: column;
    gap: 10px;
    background: white;
    padding: 10px;
    border-radius: 10px;
    box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    margin-bottom: 10px;
}}
.floating-container:hover .popup-icons {{
    display: flex;
}}
.icon-link img {{
    width: 40px;
    height: 40px;
    border-radius: 5px;
    transition: transform 0.2s;
}}
.icon-link img:hover {{
    transform: scale(1.1);
}}
</style>

<div class="floating-container">
    <div class="help-button"></div>
    <div class="popup-icons">
        <a href="https://gemini.google.com/" target="_blank" class="icon-link" title="Ir a Gemini">
            <img src="data:image/png;base64,{gemini_img_b64}" alt="Gemini">
        </a>
        <a href="https://www.deepseek.com/" target="_blank" class="icon-link" title="Ir a DeepSeek">
            <img src="data:image/png;base64,{deepseek_img_b64}" alt="DeepSeek">
        </a>
    </div>
</div>
""", unsafe_allow_html=True)

# --- Sidebar Actions ---
with st.sidebar:
    st.header("Acciones")

    # User ID / Session Control
    current_uid = get_user_id()
    new_user_input = st.text_input(
        "🆔 Usuario / Clave de Sesión",
        value=current_uid,
        help="Introduce tu nombre de usuario o clave personal para guardar y recuperar tu planeación en cualquier momento o dispositivo."
    )
    clean_uid = re.sub(r'[^a-zA-Z0-9_-]', '', new_user_input.strip())
    if clean_uid and clean_uid != current_uid:
        st.session_state["user_id"] = clean_uid
        try:
            st.query_params["user"] = clean_uid
        except Exception:
            pass
        st.session_state["_autosave_loaded"] = False
        try:
            st.rerun()
        except Exception:
            pass

    st.caption(f"🟢 Sesión activa: **{current_uid}**")
    st.checkbox("Autoguardar automáticamente", value=True, key="autosave_enabled")

    if st.button("➕ Nueva Planeación"):
        try:
            delete_user_planeacion(get_user_id())
        except Exception:
            pass
        new_uid = f"user_{uuid.uuid4().hex[:10]}"
        st.session_state["user_id"] = new_uid
        try:
            st.query_params["user"] = new_uid
        except Exception:
            pass
        for key in list(st.session_state.keys()):
            if key not in ("user_id", "autosave_enabled"):
                del st.session_state[key]
        init_session_state()
        st.session_state.quill_key_suffix += 1
        st.success("Nueva planeación iniciada.")
        try:
            st.rerun()
        except Exception:
            pass
    
    uploaded_file = st.file_uploader("Cargar Planeación (JSON)", type="json")
    
    if uploaded_file is not None:
        file_id = f"{uploaded_file.name}_{uploaded_file.size}"
        
        if file_id != st.session_state.last_loaded_file_id:
            try:
                data = json.load(uploaded_file)
                docente = data.get("docente", {})
                st.session_state.docente_titulo = docente.get("titulo", "Dr.")
                st.session_state.docente_nombre = docente.get("nombre", "")
                
                curso = data.get("curso", {})
                st.session_state.curso_grado = curso.get("grado", "1ro")
                st.session_state.curso_grupos = curso.get("grupos", [])
                st.session_state.curso_materia = curso.get("materia", "Matematicas")
                st.session_state.curso_campo = curso.get("campo", "Lenguajes")
                
                plan = data.get("planeacion", {})
                st.session_state.plan_metodologia = plan.get("metodologia", "Seleccione metodología")
                
                # Use new parse_date function
                st.session_state.plan_fecha_inicio = parse_date(plan.get("fecha_inicio"))
                st.session_state.plan_fecha_fin = parse_date(plan.get("fecha_fin"))
                
                st.session_state.plan_dias = plan.get("dias_planeados", [])
                st.session_state.plan_eje1 = plan.get("eje1", "Seleccione eje")
                st.session_state.plan_eje2 = plan.get("eje2", "Seleccione eje")
                st.session_state.plan_eje3 = plan.get("eje3", "Seleccione eje")
                st.session_state.plan_disc1 = plan.get("disciplina1", "Seleccione materia")
                st.session_state.plan_disc2 = plan.get("disciplina2", "Seleccione materia")
                st.session_state.plan_disc3 = plan.get("disciplina3", "Seleccione materia")
                
                st.session_state.text_problematica = plan.get("problematica", "")
                st.session_state.text_pda = plan.get("pda", "")
                st.session_state.pda_custom_active = plan.get("pda_custom_active", False)
                st.session_state.pda_selected = plan.get("pda_selected", "")
                st.session_state.pda_custom = plan.get("pda_custom", "")
                st.session_state.pda_entries = plan.get("pda_entries", [])
                
                # Restore official vs custom PDAs with backward compatibility
                off_sel = list(plan.get("pda_official_selected", []))
                cust_list = list(plan.get("pda_custom_entries", []))
                if not off_sel and not cust_list and plan.get("pda_entries"):
                    pdas_list = get_pdas_for_selection(st.session_state.curso_materia, st.session_state.curso_grado)
                    official_options = [p["pda"] for p in pdas_list] if pdas_list else []
                    for entry in plan.get("pda_entries", []):
                        txt = str(entry).strip()
                        if official_options and txt in official_options:
                            if txt not in off_sel: off_sel.append(txt)
                        else:
                            if txt not in cust_list: cust_list.append(txt)

                if plan.get("pda_selected") and plan.get("pda_selected") not in off_sel:
                    off_sel.append(plan.get("pda_selected"))
                if plan.get("pda_custom") and plan.get("pda_custom") not in cust_list:
                    cust_list.append(plan.get("pda_custom"))

                st.session_state.pda_official_selected = off_sel
                st.session_state.pda_custom_entries = cust_list

                st.session_state.text_objetivos = plan.get("objetivos", "")
                st.session_state.text_perfiles = plan.get("perfiles", "")
                st.session_state.text_producto = plan.get("producto", "")
                
                abpj = plan.get("secuencia_abpj", {})
                st.session_state.abpj_presentacion = abpj.get("presentacion", "")
                st.session_state.abpj_recoleccion = abpj.get("recoleccion", "")
                st.session_state.abpj_formulacion = abpj.get("formulacion", "")
                st.session_state.abpj_organizacion = abpj.get("organizacion", "")
                st.session_state.abpj_experiencia = abpj.get("experiencia", "")
                st.session_state.abpj_resultados = abpj.get("resultados", "")
                st.session_state.abpj_materiales = abpj.get("materiales", "")
                st.session_state.abpj_evaluacion = abpj.get("evaluacion", "")
                
                daily_list = plan.get("secuencia_diaria", [])
                st.session_state.daily_plan_data = {item["dia_nombre"]: item for item in daily_list}
                
                st.session_state.last_loaded_file_id = file_id
                st.session_state.quill_key_suffix += 1 # Force Quill refresh
                st.success("Planeación cargada correctamente.")
                safe_rerun()
                st.stop()
            except Exception as e:
                st.error(f"Error al cargar: {e}")

    # Attempt automatic restore from SQLite database if available and not already loaded
    def restore_autosave():
        try:
            if st.session_state.get("_autosave_loaded"):
                return
            if not st.session_state.get("autosave_enabled", True):
                return
            user_id = get_user_id()
            data = load_user_planeacion(user_id)
            if not data:
                return

            docente = data.get("docente", {})
            st.session_state.docente_titulo = docente.get("titulo", st.session_state.docente_titulo)
            st.session_state.docente_nombre = docente.get("nombre", st.session_state.docente_nombre)

            curso = data.get("curso", {})
            st.session_state.curso_grado = curso.get("grado", st.session_state.curso_grado)
            st.session_state.curso_grupos = curso.get("grupos", st.session_state.curso_grupos)
            st.session_state.curso_materia = curso.get("materia", st.session_state.curso_materia)
            st.session_state.curso_campo = curso.get("campo", st.session_state.curso_campo)

            plan = data.get("planeacion", {})
            st.session_state.plan_metodologia = plan.get("metodologia", st.session_state.plan_metodologia)
            st.session_state.plan_fecha_inicio = parse_date(plan.get("fecha_inicio"))
            st.session_state.plan_fecha_fin = parse_date(plan.get("fecha_fin"))
            st.session_state.plan_dias = plan.get("dias_planeados", st.session_state.plan_dias)
            st.session_state.plan_eje1 = plan.get("eje1", st.session_state.plan_eje1)
            st.session_state.plan_eje2 = plan.get("eje2", st.session_state.plan_eje2)
            st.session_state.plan_eje3 = plan.get("eje3", st.session_state.plan_eje3)
            st.session_state.plan_disc1 = plan.get("disciplina1", st.session_state.plan_disc1)
            st.session_state.plan_disc2 = plan.get("disciplina2", st.session_state.plan_disc2)
            st.session_state.plan_disc3 = plan.get("disciplina3", st.session_state.plan_disc3)

            st.session_state.text_problematica = plan.get("problematica", st.session_state.text_problematica)
            st.session_state.text_pda = plan.get("pda", st.session_state.text_pda)
            st.session_state.pda_custom_active = plan.get("pda_custom_active", st.session_state.pda_custom_active)
            st.session_state.pda_selected = plan.get("pda_selected", st.session_state.pda_selected)
            st.session_state.pda_custom = plan.get("pda_custom", st.session_state.pda_custom)
            st.session_state.pda_entries = plan.get("pda_entries", st.session_state.get("pda_entries", []))

            # Restore official vs custom PDAs with backward compatibility
            off_sel = list(plan.get("pda_official_selected", []))
            cust_list = list(plan.get("pda_custom_entries", []))
            if not off_sel and not cust_list and plan.get("pda_entries"):
                pdas_list = get_pdas_for_selection(st.session_state.curso_materia, st.session_state.curso_grado)
                official_options = [p["pda"] for p in pdas_list] if pdas_list else []
                for entry in plan.get("pda_entries", []):
                    txt = str(entry).strip()
                    if official_options and txt in official_options:
                        if txt not in off_sel: off_sel.append(txt)
                    else:
                        if txt not in cust_list: cust_list.append(txt)

            if plan.get("pda_selected") and plan.get("pda_selected") not in off_sel:
                off_sel.append(plan.get("pda_selected"))
            if plan.get("pda_custom") and plan.get("pda_custom") not in cust_list:
                cust_list.append(plan.get("pda_custom"))

            st.session_state.pda_official_selected = off_sel
            st.session_state.pda_custom_entries = cust_list
            st.session_state.text_objetivos = plan.get("objetivos", st.session_state.text_objetivos)
            st.session_state.text_perfiles = plan.get("perfiles", st.session_state.text_perfiles)
            st.session_state.text_producto = plan.get("producto", st.session_state.text_producto)

            abpj = plan.get("secuencia_abpj", {})
            st.session_state.abpj_presentacion = abpj.get("presentacion", st.session_state.abpj_presentacion)
            st.session_state.abpj_recoleccion = abpj.get("recoleccion", st.session_state.abpj_recoleccion)
            st.session_state.abpj_formulacion = abpj.get("formulacion", st.session_state.abpj_formulacion)
            st.session_state.abpj_organizacion = abpj.get("organizacion", st.session_state.abpj_organizacion)
            st.session_state.abpj_experiencia = abpj.get("experiencia", st.session_state.abpj_experiencia)
            st.session_state.abpj_resultados = abpj.get("resultados", st.session_state.abpj_resultados)
            st.session_state.abpj_materiales = abpj.get("materiales", st.session_state.abpj_materiales)
            st.session_state.abpj_evaluacion = abpj.get("evaluacion", st.session_state.abpj_evaluacion)
            if abpj.get("rubrica_paths"):
                st.session_state.abpj_rubrica_paths = abpj.get("rubrica_paths", st.session_state.abpj_rubrica_paths)
            else:
                legacy = abpj.get("rubrica_path")
                if legacy:
                    st.session_state.abpj_rubrica_paths = [legacy]

            daily_list = plan.get("secuencia_diaria", [])
            if isinstance(daily_list, list):
                st.session_state.daily_plan_data = {item["dia_nombre"]: item for item in daily_list if "dia_nombre" in item}

            st.session_state.last_loaded_file_id = "autosave"
            st.session_state.quill_key_suffix += 1
            st.session_state["_autosave_loaded"] = True
        except Exception as e:
            st.warning(f"No se pudo restaurar autosave: {e}")

    restore_autosave()
    cleanup_temp_rubric_files()


    # Login/AppsScript integration removed — app now shows planner directly

    def get_current_data():
        daily_sequence = []
        start = st.session_state.plan_fecha_inicio
        end = st.session_state.plan_fecha_fin
        active_days = st.session_state.plan_dias
        dias_map = {0: "Lunes", 1: "Martes", 2: "Miércoles", 3: "Jueves", 4: "Viernes"}
        
        if start <= end:
            delta = end - start
            for i in range(delta.days + 1):
                current = start + timedelta(days=i)
                wd = current.weekday()
                if wd in dias_map:
                    dia_nombre = dias_map[wd]
                    if dia_nombre in active_days:
                        key = f"{dia_nombre} {current.strftime('%d/%m/%Y')}"
                        if key in st.session_state.daily_plan_data:
                            daily_sequence.append(st.session_state.daily_plan_data[key])
                        else:
                            daily_sequence.append({
                                "dia_nombre": key, "inicio": "", "desarrollo": "", "cierre": "",
                                    "materiales": "", "evaluacion": "", "rubrica_paths": []
                            })

        return {
            "docente": {
                "titulo": st.session_state.docente_titulo,
                "nombre": st.session_state.docente_nombre
            },
            "curso": {
                "grado": st.session_state.curso_grado,
                "grupos": st.session_state.curso_grupos,
                "materia": st.session_state.curso_materia,
                "campo": st.session_state.curso_campo
            },
            "planeacion": {
                "metodologia": st.session_state.plan_metodologia,
                "fecha_inicio": st.session_state.plan_fecha_inicio.isoformat(),
                "fecha_fin": st.session_state.plan_fecha_fin.isoformat(),
                "dias_planeados": st.session_state.plan_dias,
                "problematica": st.session_state.text_problematica,
                "pda": get_current_pda(),
                "pda_entries": st.session_state.get("pda_entries", []),
                "pda_custom_active": st.session_state.get("pda_custom_active", False),
                "pda_selected": st.session_state.get("pda_selected", ""),
                "pda_custom": st.session_state.get("pda_custom", ""),
                "objetivos": st.session_state.text_objetivos,
                "perfiles": st.session_state.text_perfiles,
                "producto": st.session_state.text_producto,
                "eje1": st.session_state.plan_eje1,
                "eje2": st.session_state.plan_eje2,
                "eje3": st.session_state.plan_eje3,
                "disciplina1": st.session_state.plan_disc1,
                "disciplina2": st.session_state.plan_disc2,
                "disciplina3": st.session_state.plan_disc3,
                "secuencia_abpj": {
                    "presentacion": st.session_state.abpj_presentacion,
                    "recoleccion": st.session_state.abpj_recoleccion,
                    "formulacion": st.session_state.abpj_formulacion,
                    "organizacion": st.session_state.abpj_organizacion,
                    "experiencia": st.session_state.abpj_experiencia,
                    "resultados": st.session_state.abpj_resultados,
                    "materiales": st.session_state.abpj_materiales,
                    "evaluacion": st.session_state.abpj_evaluacion,
                    "rubrica_paths": st.session_state.get("abpj_rubrica_paths", [])
                },
                "secuencia_diaria": daily_sequence
            }
        }

    json_data = json.dumps(get_current_data(), indent=4, ensure_ascii=False)
    st.download_button("Guardar Planeación (JSON)", data=json_data, file_name="planeacion.json", mime="application/json")


# --- Main UI ---

col1, col2, col3 = st.columns([1, 3, 1])
with col1:
    logo_imm = _get_full_path("LOGO imm.png")
    if os.path.exists(logo_imm):
        st.image(logo_imm, width=150)
with col2:
    st.markdown("""
    <div style='text-align: center;'>
        <h3>Secretaría De Educación Pública</h3>
        <h4>Dirección De Educación Secundaria</h4>
        <h4>Instituto Mexicano Madero</h4>
        <h2>Planeaciones Docente</h2>
    </div>
    """, unsafe_allow_html=True)
with col3:
    logo_sep = _get_full_path("logo_sep.png")
    if os.path.exists(logo_sep):
        st.image(logo_sep, width=180)

st.markdown("---")

tab1, tab2, tab3, tab4 = st.tabs(["Docente y Curso", "Detalles Generales", "Contenido", "Secuencia Didáctica"])

with tab1:
    col_d1, col_d2 = st.columns(2)
    with col_d1:
        st.subheader("Información Docente")
        st.selectbox("Título", ["Dr.", "Dra.", "Mtro.", "Mtra.", "Prof.", "Pasante."], key="docente_titulo")
        st.text_input("Nombre Completo", key="docente_nombre")
    
    with col_d2:
        st.subheader("Información Curso")
        c1, c2 = st.columns(2)
        with c1:
            st.selectbox("Grado", ["1ro", "2do", "3ro"], key="curso_grado")
        with c2:
                st.multiselect("Grupos", LISTA_GRUPOS, key="curso_grupos")
                st.checkbox("Bilingüe", key="curso_bilingue")
        
        st.selectbox("Materia", LISTA_MATERIAS, key="curso_materia")
        st.selectbox("Campo Formativo", LISTA_CAMPOS, key="curso_campo")

with tab2:
    st.subheader("Detalles de la Planeación")
    st.selectbox("Metodología", LISTA_METODOLOGIA, key="plan_metodologia")
    
    col_t1, col_t2 = st.columns(2)
    with col_t1:
        st.date_input("Fecha Inicio", format="DD/MM/YYYY", key="plan_fecha_inicio")
    with col_t2:
        st.date_input("Fecha Fin", format="DD/MM/YYYY", key="plan_fecha_fin")
    
    st.multiselect("Días de Clase", LISTA_DIAS, key="plan_dias")
    
    st.markdown("**Ejes Articuladores**")
    c_e1, c_e2, c_e3 = st.columns(3)
    with c_e1: st.selectbox("Eje 1", LISTA_EJES, key="plan_eje1")
    with c_e2: st.selectbox("Eje 2", LISTA_EJES, key="plan_eje2")
    with c_e3: st.selectbox("Eje 3", LISTA_EJES, key="plan_eje3")
    
    st.markdown("**Materias Vinculadas**")
    c_m1, c_m2, c_m3 = st.columns(3)
    with c_m1: st.selectbox("Materia 1", ["Seleccione materia"] + LISTA_MATERIAS, key="plan_disc1")
    with c_m2: st.selectbox("Materia 2", ["Seleccione materia"] + LISTA_MATERIAS, key="plan_disc2")
    with c_m3: st.selectbox("Materia 3", ["Seleccione materia"] + LISTA_MATERIAS, key="plan_disc3")

with tab3:
    st.subheader("Contenido Pedagógico")
    # Quill Editor Configuration
    toolbar = [
        ['bold', 'italic', 'underline'],
        [{'list': 'ordered'}, {'list': 'bullet'}]
    ]
    
    # Use key suffix to force refresh
    ks = st.session_state.quill_key_suffix
    
    st.session_state.text_problematica = st_quill(value=st.session_state.text_problematica, placeholder="Problemática Contextual", toolbar=toolbar, key=f"quill_prob_{ks}")
    
    st.markdown("### Proceso de Desarrollo de Aprendizaje (PDA)")
    pdas_list = get_pdas_for_selection(st.session_state.curso_materia, st.session_state.curso_grado)
    official_options = [p["pda"] for p in pdas_list] if pdas_list else []
    
    st.markdown("#### 1. PDAs Oficiales (Plan Sintético)")
    if official_options:
        st.multiselect(
            "Seleccione uno o varios PDAs oficiales:",
            options=official_options,
            key="pda_official_selected",
            help="Puede seleccionar múltiples PDAs del Plan Sintético."
        )
        if st.session_state.get("pda_official_selected"):
            with st.expander("Ver contenidos del Plan Sintético asociados", expanded=False):
                for sel in st.session_state.pda_official_selected:
                    rec = next((r for r in pdas_list if r["pda"] == sel), None)
                    if rec:
                        st.caption(f"• **PDA:** {sel}<br/>  **Contenido:** {rec['contenido']}", unsafe_allow_html=True)
    else:
        st.info("No se encontraron PDAs oficiales para esta materia y grado en el Plan Sintético.")
    
    st.markdown("#### 2. PDAs Personalizados (Codiseño)")
    if "pda_custom_entries" not in st.session_state or not isinstance(st.session_state.pda_custom_entries, list):
        st.session_state.pda_custom_entries = []

    new_custom_entries = []
    for idx, entry in enumerate(st.session_state.get("pda_custom_entries", [])):
        c_pda1, c_pda2 = st.columns([5, 1])
        with c_pda1:
            val = st_quill(
                value=entry,
                placeholder=f"Redactar PDA personalizado {idx+1}",
                toolbar=toolbar,
                key=f"quill_pda_custom_{ks}_{idx}"
            )
        with c_pda2:
            if st.button(f"🗑️ Eliminar", key=f"del_pda_custom_{idx}"):
                continue
        new_custom_entries.append(val)

    if st.button("➕ Agregar PDA Personalizado (Codiseño)"):
        new_custom_entries.append("")
        st.session_state.pda_custom_entries = new_custom_entries
        safe_rerun()

    st.session_state.pda_custom_entries = new_custom_entries

    st.session_state.text_objetivos = st_quill(value=st.session_state.text_objetivos, placeholder="Objetivos", toolbar=toolbar, key=f"quill_obj_{ks}")
    st.session_state.text_perfiles = st_quill(value=st.session_state.text_perfiles, placeholder="Perfiles de Egreso", toolbar=toolbar, key=f"quill_perf_{ks}")
    st.session_state.text_producto = st_quill(value=st.session_state.text_producto, placeholder="Producto Final", toolbar=toolbar, key=f"quill_prod_{ks}")

with tab4:
    st.subheader("Secuencia Didáctica")
    toolbar_simple = [['bold', 'italic', 'underline'], [{'list': 'bullet'}]]
    ks = st.session_state.quill_key_suffix
    
    if "ABPj" in st.session_state.plan_metodologia:
        st.markdown("### Aprendizaje Basado en Proyectos (ABPj)")
        
        col_abp1, col_abp2 = st.columns(2)
        with col_abp1:
            st.session_state.abpj_presentacion = st_quill(value=st.session_state.abpj_presentacion, placeholder="1. Presentación", toolbar=toolbar_simple, key=f"q_abpj_1_{ks}")
            st.session_state.abpj_formulacion = st_quill(value=st.session_state.abpj_formulacion, placeholder="3. Formulación del Problema", toolbar=toolbar_simple, key=f"q_abpj_3_{ks}")
            st.session_state.abpj_experiencia = st_quill(value=st.session_state.abpj_experiencia, placeholder="5. Vivamos la Experiencia", toolbar=toolbar_simple, key=f"q_abpj_5_{ks}")
            st.session_state.abpj_materiales = st_quill(value=st.session_state.abpj_materiales, placeholder="Materiales", toolbar=toolbar_simple, key=f"q_abpj_mat_{ks}")
        
        with col_abp2:
            st.session_state.abpj_recoleccion = st_quill(value=st.session_state.abpj_recoleccion, placeholder="2. Recolección", toolbar=toolbar_simple, key=f"q_abpj_2_{ks}")
            st.session_state.abpj_organizacion = st_quill(value=st.session_state.abpj_organizacion, placeholder="4. Organización del Proyecto", toolbar=toolbar_simple, key=f"q_abpj_4_{ks}")
            st.session_state.abpj_resultados = st_quill(value=st.session_state.abpj_resultados, placeholder="6. Resultados y Análisis", toolbar=toolbar_simple, key=f"q_abpj_6_{ks}")
            
            st.markdown("#### Evaluación")
            st.session_state.abpj_evaluacion = st_quill(value=st.session_state.abpj_evaluacion, placeholder="Evaluación", toolbar=toolbar_simple, key=f"q_abpj_eval_{ks}")
            uploaded_rubric = st.file_uploader("Anexar Rúbrica (Imagen)", type=["png", "jpg", "jpeg"], key=f"abpj_rubric_uploader_{ks}")
            if uploaded_rubric:
                st.session_state.abpj_rubrica_paths = process_uploaded_rubrics(
                    uploaded_rubric,
                    st.session_state.get("abpj_rubrica_paths", []),
                    "abpj"
                )
            
            if st.session_state.get("abpj_rubrica_paths"):
                st.caption(f"Rúbricas actuales: {', '.join([os.path.basename(x) for x in st.session_state.abpj_rubrica_paths])}")
                for idx, path in enumerate(st.session_state.abpj_rubrica_paths.copy()):
                    if os.path.exists(path):
                        col_a, col_b = st.columns([4, 1])
                        with col_a:
                            st.image(path, width=120, caption=os.path.basename(path))
                        with col_b:
                            if st.button(f"Eliminar {os.path.basename(path)}", key=f"del_abpj_rubrica_{idx}"):
                                remove_rubric_image(path, st.session_state.abpj_rubrica_paths)
                    else:
                        st.write(f"Ruta no encontrada: {path}")

    elif st.session_state.plan_metodologia != "Seleccione metodología":
        st.markdown(f"### Planeación Diaria ({st.session_state.plan_metodologia})")
        
        start = st.session_state.plan_fecha_inicio
        end = st.session_state.plan_fecha_fin
        active_days = st.session_state.plan_dias
        dias_map = {0: "Lunes", 1: "Martes", 2: "Miércoles", 3: "Jueves", 4: "Viernes"}
        
        dias_generados = []
        if start <= end:
            delta = end - start
            for i in range(delta.days + 1):
                current = start + timedelta(days=i)
                wd = current.weekday()
                if wd in dias_map:
                    dia_nombre = dias_map[wd]
                    if dia_nombre in active_days:
                        dias_generados.append((current, dia_nombre))
        
        if not dias_generados:
            st.warning("No hay días hábiles seleccionados en el rango de fechas.")
        else:
            for i, (fecha_obj, dia_nombre) in enumerate(dias_generados):
                fecha_str = fecha_obj.strftime("%d/%m/%Y")
                key_base = f"{dia_nombre} {fecha_str}"
                
                if key_base not in st.session_state.daily_plan_data:
                    st.session_state.daily_plan_data[key_base] = {
                        "dia_nombre": key_base, "inicio": "", "desarrollo": "", "cierre": "",
                        "materiales": "", "evaluacion": "", "rubrica_paths": []
                    }
                
                day_data = st.session_state.daily_plan_data[key_base]
                
                with st.expander(f"Sesión {i+1}: {key_base}", expanded=True):
                    c1, c2, c3 = st.columns(3)
                    with c1: day_data["inicio"] = st_quill(value=day_data["inicio"], placeholder="Inicio", key=f"inicio_{key_base}_{ks}", toolbar=toolbar_simple)
                    with c2: day_data["desarrollo"] = st_quill(value=day_data["desarrollo"], placeholder="Desarrollo", key=f"desarrollo_{key_base}_{ks}", toolbar=toolbar_simple)
                    with c3: day_data["cierre"] = st_quill(value=day_data["cierre"], placeholder="Cierre", key=f"cierre_{key_base}_{ks}", toolbar=toolbar_simple)
                    
                    c4, c5 = st.columns(2)
                    with c4: day_data["materiales"] = st_quill(value=day_data["materiales"], placeholder="Materiales", key=f"mat_{key_base}_{ks}", toolbar=toolbar_simple)
                    with c5: 
                        day_data["evaluacion"] = st_quill(value=day_data["evaluacion"], placeholder="Evaluación", key=f"eval_{key_base}_{ks}", toolbar=toolbar_simple)
                        u_rubrics = st.file_uploader("Rúbrica (puede subir varias)", type=["png", "jpg", "jpeg"], key=f"up_{key_base}_{ks}", accept_multiple_files=True)
                        if u_rubrics:
                            day_data["rubrica_paths"] = process_uploaded_rubrics(
                                u_rubrics,
                                day_data.get("rubrica_paths", []),
                                f"day_{i}"
                            )

                        if day_data.get("rubrica_paths"):
                            for idx, path in enumerate(day_data["rubrica_paths"].copy()):
                                if os.path.exists(path):
                                    col_a, col_b = st.columns([4, 1])
                                    with col_a:
                                        st.image(path, width=120, caption=os.path.basename(path))
                                    with col_b:
                                        if st.button(f"Eliminar {os.path.basename(path)}", key=f"del_day_rubrica_{key_base}_{idx}"):
                                            remove_rubric_image(path, day_data["rubrica_paths"])
                                else:
                                    st.write(f"Ruta no encontrada: {path}")

# --- AI Prompt Generation ---
st.markdown("---")
autosave_current_data()
if st.button("✨ Generar Prompt IA"):
    d = get_current_data()
    p_data = d['planeacion']
    
    grado = d['curso']['grado']
    edad = "11 a 12 años" if "1" in grado else "12 a 13 años" if "2" in grado else "13 a 15 años"
    
    ejes = ", ".join([e for e in [p_data['eje1'], p_data['eje2'], p_data['eje3']] if e and "Seleccione" not in e])
    disc = ", ".join([x for x in [p_data['disciplina1'], p_data['disciplina2'], p_data['disciplina3']] if x and "Seleccione" not in x])
    dias_txt = ", ".join(p_data['dias_planeados'])
    
    # Clean HTML for prompt (basic strip)
    def clean_html(t): return re.sub(r'<[^>]+>', '', t) if t else ""
    
    prompt = f"Actúa como un docente experto de secundaria en México (SEP, Nueva Escuela Mexicana).\n\n"
    prompt += f"Genera una planeación didáctica para la materia de **{d['curso']['materia']}**.\n"
    prompt += f"- **Grado:** {grado} (Alumnos de aprox. {edad}).\n"
    prompt += f"- **Campo Formativo:** {d['curso']['campo']}.\n"
    prompt += f"- **Metodología:** {p_data['metodologia']}.\n"
    prompt += f"- **Temporalidad:** Del {p_data['fecha_inicio']} al {p_data['fecha_fin']}.\n"
    prompt += f"- **Días de clase:** {dias_txt}.\n"
    if ejes: prompt += f"- **Ejes Articuladores:** {ejes}.\n"
    if disc: prompt += f"- **Materias vinculadas:** {disc}.\n"
    
    prompt += f"- **Problemática Contextual:** {clean_html(p_data['problematica']) or 'No definida. Propón una relevante.'}\n"
    if p_data.get('pda_custom_active', False):
        prompt += f"- **PDA redactado como codiseño:** {clean_html(p_data.get('pda_custom', '')) or 'No definido.'}\n"
        prompt += f"- **Instrucción Especial:** Redacta un PDA que se ajuste a la SEP fase 6.\n"
    else:
        prompt += f"- **PDA (Plan Sintético):** {clean_html(p_data.get('pda_selected', '')) or 'Propón el PDA oficial más adecuado.'}\n"
    prompt += f"- **Objetivos:** {clean_html(p_data['objetivos']) or 'Propón objetivos de aprendizaje adecuados.'}\n"
    prompt += f"- **Perfil de Egreso:** {clean_html(p_data['perfiles']) or 'Propón los rasgos del perfil de egreso que se favorecen.'}\n"
    prompt += f"- **Producto Final:** {clean_html(p_data['producto']) or 'Propón un producto creativo.'}\n"
    
    if "Proyectos" in p_data['metodologia']:
        prompt += "\nDesarrolla la secuencia didáctica siguiendo las fases de **ABPj** (Presentación, Recolección, Formulación, Organización, Vivamos la experiencia, Resultados)."
    else:
        prompt += "\nDesarrolla la secuencia didáctica como una **Planeación Diaria** (Inicio, Desarrollo, Cierre) para cada día."

    prompt += "\nEn la secuencia didáctica, incluye para cada actividad/fase los **Materiales y Recursos** necesarios, así como una propuesta de **Evaluación** (sugiere Lista de Cotejo o Rúbrica si es necesario)."

    # If bilingual option is active, append instruction within the single prompt
    if st.session_state.get("curso_bilingue"):
        materia_bilingue = d['curso']['materia']
        prompt += f"\n\n**INSTRUCCIÓN BILINGÜE:** Genera la planeación didáctica completamente en inglés para la materia de {materia_bilingue}. Todos los títulos, descripciones, actividades y contenidos deben estar escritos en inglés."

    st.code(prompt, language="text")
    st.info("Copia el texto de arriba y pégalo en tu IA favorita (ChatGPT, Gemini, DeepSeek).")

# --- PDF Generation ---
def generate_pdf_bytes():
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter), topMargin=0.5*inch, bottomMargin=0.5*inch, leftMargin=0.5*inch, rightMargin=0.5*inch)
    elements = []
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle('Body', parent=styles['Normal'], spaceAfter=6, leading=14)
    PB = lambda x: Paragraph(f"<b>{x}</b>", body_style)
    P = lambda x: Paragraph(html_to_reportlab(str(x)), body_style)
    
    ruta_logo_imm = _get_full_path("LOGO imm.png")
    ruta_logo_sep = _get_full_path("logo_sep.png")
    
    if os.path.exists(ruta_logo_imm):
        logo_imm_rl = RLImage(ruta_logo_imm, width=1.5*inch, height=0.75*inch, kind='proportional')
    else:
        logo_imm_rl = Paragraph("[LOGO IMM]", styles['Normal'])
        
    if os.path.exists(ruta_logo_sep):
        logo_sep_rl = RLImage(ruta_logo_sep, width=1.8*inch, height=0.75*inch, kind='proportional')
    else:
        logo_sep_rl = Paragraph("[LOGO SEP]", styles['Normal'])

    header_paragraphs = [Paragraph(t, ParagraphStyle(name='HeaderCenter', alignment=TA_CENTER, fontName='Helvetica-Bold', fontSize=11, leading=14)) for t in ["Secretaría De Educación Pública", "Dirección De Educación Secundaria", "Instituto Mexicano Madero", "Planeaciones Docente"]]
    
    page_width = landscape(letter)[0] - 1*inch
    col_widths_header = [2*inch, page_width - 4*inch, 2*inch]
    header_data = [[logo_imm_rl, header_paragraphs, logo_sep_rl]]
    header_table = Table(header_data, colWidths=col_widths_header)
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
        ('ALIGN', (2, 0), (2, 0), 'RIGHT'),
        ('BACKGROUND', (0,0), (-1,-1), colors.white),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6)
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 0.2 * inch))

    d = get_current_data()
    p = d['planeacion']
    
    # PB and P are now defined above with body_style for consistent paragraph spacing
    grupos_str = ", ".join(d['curso']['grupos'])
    docente_name = f"{d['docente']['titulo']} {d['docente']['nombre']}"
    dias_str = ", ".join(p['dias_planeados'])
    
    # Format dates for PDF
    f_inicio = date.fromisoformat(p['fecha_inicio']).strftime("%d/%m/%Y")
    f_fin = date.fromisoformat(p['fecha_fin']).strftime("%d/%m/%Y")
    
    temp_str = f"Del {f_inicio} al {f_fin}. Días: {dias_str}"
    
    ejes = ", ".join([e for e in [p['eje1'], p['eje2'], p['eje3']] if e and "Seleccione" not in e])
    disc = ", ".join([x for x in [p['disciplina1'], p['disciplina2'], p['disciplina3']] if x and "Seleccione" not in x])

    row0 = [[PB("Escuela:"), P("Instituto Mexicano Madero"), PB("CCT:"), P("21PES0013L"), PB("Docente:"), P(docente_name)]]
    t0 = Table(row0, colWidths=[0.8*inch, 2.7*inch, 0.5*inch, 1*inch, 0.8*inch, 4.2*inch])
    t0.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BACKGROUND', (0,0), (-1,-1), colors.white),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
        ('RIGHTPADDING', (0,0), (-1,-1), 4)
    ]))
    
    row1 = [[PB("Grado:"), P(d['curso']['grado']), PB("Grupo:"), P(grupos_str), PB("Fase:"), P("6"), PB("Campo:"), P(d['curso']['campo'])]]
    t1 = Table(row1, colWidths=[0.8*inch, 1.2*inch, 0.8*inch, 1.2*inch, 0.8*inch, 0.8*inch, 1.4*inch, 3*inch])
    t1.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BACKGROUND', (0,0), (-1,-1), colors.white),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
        ('RIGHTPADDING', (0,0), (-1,-1), 4)
    ]))

    elements.extend([t0, Spacer(1, 0.1 * inch), t1, Spacer(1, 0.2 * inch)])

    # PDA label and optional legend: use the formal name for printouts
    if p.get("pda_custom_active", False):
        pda_label = "PDA redactado como codiseño:"
        pda_legend = None
    else:
        pda_label = "Proceso de Desarrollo de Aprendizaje (PDA):"
        pda_legend = "Plan Sintético (oficial)" if p.get("pda_selected", "") else None

    # Build main data rows with safe splitting: if content is very large, place it outside the table
    rows = [
        ("Materia:", p['materia'] if 'materia' in p else d['curso']['materia']),
        ("Metodología:", p['metodologia']),
        ("Ejes:", ejes),
        ("Vinculación:", disc),
        ("Problemática:", p['problematica']),
        (pda_label, p.get('pda') or (p.get('pda_selected') or p.get('pda_custom', ''))),
        ("Objetivos:", p['objetivos']),
        ("Perfiles:", p['perfiles']),
        ("Temporalidad:", temp_str),
        ("Producto:", p['producto'])
    ]

    small_table_rows = []
    for label, content_html in rows:
        plain = re.sub(r'<[^>]+>', '', str(content_html or ''))
        if len(plain) > 900 or plain.count('\n') > 20:
            # Append as full-width block
            elements.append(PB(label))
            flowables = html_to_flowables(content_html, styles)
            for f in flowables:
                elements.append(f)
        else:
            # small enough to include in table
            cell_flow = html_to_reportlab(str(content_html))
            small_table_rows.append([PB(label), Paragraph(cell_flow, styles['Normal'])])

    if small_table_rows:
        main_table = Table(small_table_rows, colWidths=[2.0*inch, page_width - 2.0*inch])
        main_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BACKGROUND', (0,0), (-1,-1), colors.white),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4)
        ]))
        elements.append(main_table)

    if "ABPj" in p['metodologia']:
        elements.append(PageBreak())
        elements.append(Paragraph("Secuencia Didáctica (ABPj)", ParagraphStyle(name='H2', fontSize=12, fontName='Helvetica-Bold', alignment=TA_CENTER)))
        abpj = p['secuencia_abpj']
        seq_data = []
        campos = [("Presentación", "presentacion"), ("Recolección", "recoleccion"), ("Formulación", "formulacion"), ("Organización", "organizacion"), ("Vivamos", "experiencia"), ("Resultados", "resultados"), ("Materiales", "materiales")]
        for label, key in campos:
            seq_data.append([PB(label), Paragraph(html_to_reportlab(abpj.get(key, "")), styles['Normal'])])
        st_table = Table(seq_data, colWidths=[2.0*inch, page_width - 2.0*inch])
        st_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BACKGROUND', (0,0), (-1,-1), colors.white),
            ('LEFTPADDING', (0,0), (-1,-1), 6),
            ('RIGHTPADDING', (0,0), (-1,-1), 6)
        ]))
        elements.append(st_table)
        elements.append(Spacer(1, 0.1*inch))
        elements.append(PB("Evaluación"))
        eval_flow = html_to_flowables(abpj.get("evaluacion", ""), styles)
        for img_path in abpj.get("rubrica_paths", []):
            _embed_image_to_pdf(img_path, eval_flow, page_width - 2.0*inch)
        for f in eval_flow:
            elements.append(f)
    elif p['metodologia'] != "Seleccione metodología":
        elements.append(PageBreak())
        elements.append(Paragraph("Secuencia Didáctica (Diaria)", ParagraphStyle(name='H2', fontSize=12, fontName='Helvetica-Bold', alignment=TA_CENTER)))
        daily = p['secuencia_diaria']
        for i, day in enumerate(daily):
            elements.append(Paragraph(f"<b>{day['dia_nombre']}</b>", styles['Normal']))
            d_data = [
                [PB("Inicio"), Paragraph(html_to_reportlab(day.get("inicio", "")), styles['Normal'])],
                [PB("Desarrollo"), Paragraph(html_to_reportlab(day.get("desarrollo", "")), styles['Normal'])],
                [PB("Cierre"), Paragraph(html_to_reportlab(day.get("cierre", "")), styles['Normal'])],
                [PB("Materiales"), Paragraph(html_to_reportlab(day.get("materiales", "")), styles['Normal'])]
            ]
            dt = Table(d_data, colWidths=[2.0*inch, page_width - 2.0*inch])
            dt.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BACKGROUND', (0,0), (-1,-1), colors.white),
                ('LEFTPADDING', (0,0), (-1,-1), 6),
                ('RIGHTPADDING', (0,0), (-1,-1), 6)
            ]))
            elements.append(dt)
            elements.append(Spacer(1, 0.1*inch))
            elements.append(PB("Evaluación"))
            eval_flow = html_to_flowables(day.get("evaluacion", ""), styles)
            for img_path in day.get("rubrica_paths", []):
                _embed_image_to_pdf(img_path, eval_flow, page_width - 2.0*inch)
            for f in eval_flow:
                elements.append(f)
            elements.append(Spacer(1, 0.1*inch))

    elements.append(Spacer(1, 1.5*inch))
    sig_style = ParagraphStyle(name='Firma', alignment=TA_CENTER)
    sig_table = Table([[Paragraph("_____________________________________________", sig_style)], [Paragraph("Vo. Bo. Director David Pérez Ordoñez", sig_style)]], colWidths=[page_width])
    sig_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.white),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
        ('RIGHTPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6)
    ]))
    elements.append(sig_table)

    campo = d['curso']['campo']
    bg_image_name = None
    if campo == "Lenguajes":
        bg_image_name = "Lenguajes Hoja.png"
    elif campo == "Saberes y Pensamiento Científico":
        bg_image_name = "Pensamiento Cientifico hoja.png"
    elif campo == "Ética, Naturaleza y Sociedades":
        bg_image_name = "Etica hoja.png"
    elif campo == "De lo Humano y lo Comunitario":
        bg_image_name = "Comunitario hoja.png"

    def draw_background(canvas, doc):
        canvas.saveState()
        if bg_image_name:
            bg_path = _get_full_path(bg_image_name)
            if os.path.exists(bg_path):
                try:
                    flat = flatten_image(bg_path)
                    width, height = doc.pagesize
                    canvas.drawImage(flat, 0, 0, width=width, height=height)
                except Exception:
                    canvas.drawImage(bg_path, 0, 0)
        canvas.restoreState()

    doc.build(elements, onFirstPage=draw_background, onLaterPages=draw_background)
    buffer.seek(0)
    return buffer


def _add_docx_header_logos(doc):
    """Add institutional logos to the header of a Word document."""
    section = doc.sections[0]
    header = section.header
    # Clear any default paragraph text
    for para in header.paragraphs:
        para.text = ''

    # We build a table inside the header: [IMM logo] [center space] [SEP logo]
    header_table = header.add_table(rows=1, cols=3, width=Emu(int((
        section.page_width - section.left_margin - section.right_margin
    ))))
    header_table.style = 'Table Grid'
    # Remove borders from header table
    tbl = header_table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    cells = header_table.rows[0].cells

    # Left cell: IMM logo
    ruta_logo_imm = _get_full_path("LOGO imm.png")
    if os.path.exists(ruta_logo_imm):
        try:
            run_imm = cells[0].paragraphs[0].add_run()
            run_imm.add_picture(ruta_logo_imm, width=Inches(1.2))
        except Exception:
            cells[0].paragraphs[0].text = 'IMM'
    else:
        cells[0].paragraphs[0].text = 'IMM'

    # Center cell: institution name
    center_para = cells[1].paragraphs[0]
    center_para.alignment = 1  # CENTER
    run_c = center_para.add_run('Instituto Mexicano Madero')
    run_c.bold = True

    # Right cell: SEP logo
    ruta_logo_sep = _get_full_path("logo_sep.png")
    if os.path.exists(ruta_logo_sep):
        try:
            right_para = cells[2].paragraphs[0]
            right_para.alignment = 2  # RIGHT
            run_sep = right_para.add_run()
            run_sep.add_picture(ruta_logo_sep, width=Inches(1.5))
        except Exception:
            cells[2].paragraphs[0].text = 'SEP'
    else:
        cells[2].paragraphs[0].text = 'SEP'


def generate_docx_bytes():
    d = get_current_data()
    doc = Document()
    set_docx_landscape(doc)

    # Apply narrow margins (0.5 inch on all sides)
    section = doc.sections[0]
    narrow = Inches(0.5)
    section.top_margin = narrow
    section.bottom_margin = narrow
    section.left_margin = narrow
    section.right_margin = narrow

    # Add logos to header
    _add_docx_header_logos(doc)

    doc.add_heading('Planeación Docente', level=1)
    curso = d['curso']
    docente = d['docente']
    doc.add_paragraph(f"Escuela: Instituto Mexicano Madero")
    doc.add_paragraph(f"Docente: {docente.get('titulo','')} {docente.get('nombre','')}")
    doc.add_paragraph(f"Materia: {curso.get('materia','')}")
    doc.add_paragraph(f"Grado: {curso.get('grado','')} Grupos: {', '.join(curso.get('grupos',[]))}")

    p = d['planeacion']
    doc.add_heading('Contenido Pedagógico', level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(2)
    table.columns[1].width = Inches(8)

    def add_row(label, content):
        row = table.add_row().cells
        row[0].paragraphs[0].add_run(label).bold = True
        fill_docx_cell_from_html(row[1], content)

    add_row('Metodología:', p.get('metodologia', ''))
    add_row('Temporalidad:', f"{date.fromisoformat(p.get('fecha_inicio','')).strftime('%d/%m/%Y')} - {date.fromisoformat(p.get('fecha_fin','')).strftime('%d/%m/%Y')}")
    add_row('Días de clase:', ', '.join(p.get('dias_planeados', [])))
    add_row('Ejes articuladores:', ', '.join([e for e in [p.get('eje1'), p.get('eje2'), p.get('eje3')] if e and 'Seleccione' not in e]))
    add_row('Materias vinculadas:', ', '.join([x for x in [p.get('disciplina1'), p.get('disciplina2'), p.get('disciplina3')] if x and 'Seleccione' not in x]))
    add_row('Problemática Contextual:', p.get('problematica',''))
    add_row('Objetivos:', p.get('objetivos',''))
    add_row('Perfil de Egreso:', p.get('perfiles',''))
    add_row('Producto Final:', p.get('producto',''))

    pdas = d['planeacion'].get('pda_entries') or []
    if pdas:
        doc.add_heading('PDA(s)', level=3)
        for idx, pd in enumerate(pdas):
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(re.sub(r'<[^>]+>', '', str(pd)))

    if 'ABPj' in p.get('metodologia',''):
        doc.add_heading('Secuencia Didáctica ABPj', level=2)
        abpj = p.get('secuencia_abpj', {})
        abpj_table = doc.add_table(rows=0, cols=2)
        abpj_table.style = 'Table Grid'
        abpj_table.autofit = False
        abpj_table.columns[0].width = Inches(2)
        abpj_table.columns[1].width = Inches(8)
        for label in ['presentacion', 'recoleccion', 'formulacion', 'organizacion', 'experiencia', 'resultados', 'materiales', 'evaluacion']:
            label_text = label.capitalize()
            row = abpj_table.add_row().cells
            row[0].paragraphs[0].add_run(f"{label_text}:").bold = True
            fill_docx_cell_from_html(row[1], abpj.get(label, ''))
            if label == 'evaluacion':
                for img_path in abpj.get('rubrica_paths', []):
                    try:
                        doc.add_picture(img_path, width=Inches(4))
                    except Exception:
                        doc.add_paragraph(f"[Imagen: {os.path.basename(img_path)}]")
    else:
        doc.add_heading('Secuencia Didáctica', level=2)
        for day in d['planeacion'].get('secuencia_diaria', []):
            doc.add_heading(day.get('dia_nombre', ''), level=3)
            day_table = doc.add_table(rows=0, cols=2)
            day_table.style = 'Table Grid'
            day_table.autofit = False
            day_table.columns[0].width = Inches(2)
            day_table.columns[1].width = Inches(8)
            for label, content in [
                ('Inicio:', day.get('inicio','')),
                ('Desarrollo:', day.get('desarrollo','')),
                ('Cierre:', day.get('cierre','')),
                ('Materiales:', day.get('materiales','')),
                ('Evaluación:', day.get('evaluacion',''))
            ]:
                row = day_table.add_row().cells
                row[0].paragraphs[0].add_run(label).bold = True
                fill_docx_cell_from_html(row[1], content)
            for img_path in day.get('rubrica_paths', []):
                try:
                    doc.add_picture(img_path, width=Inches(4))
                except Exception:
                    doc.add_paragraph(f"[Imagen: {os.path.basename(img_path)}]")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


st.markdown("### Descargar Planeación")
download_choice = st.selectbox("Formato de descarga:", ["PDF", "WORD", "PDF + WORD"]) 
if st.button("📥 Generar y Descargar"):
    try:
        # build filename
        d = get_current_data()
        periodo = f"{date.fromisoformat(d['planeacion']['fecha_inicio']).strftime('%d %b')} - {date.fromisoformat(d['planeacion']['fecha_fin']).strftime('%d %b %Y')}"
        grado = f"{d['curso'].get('grado','')}-{' '.join(d['curso'].get('grupos',[]))}"
        docente_str = f"{d['docente'].get('titulo','')} {d['docente'].get('nombre','') }"
        base_name = f"Planeacion-{periodo}-{d['curso'].get('materia','')}-{grado}-{docente_str}"
        base_name = safe_filename(base_name)

        if download_choice in ("PDF", "PDF + WORD"):
            pdf_buf = generate_pdf_bytes()
            st.download_button(label="Descargar PDF", data=pdf_buf, file_name=base_name + ".pdf", mime="application/pdf")

        if download_choice in ("WORD", "PDF + WORD"):
            docx_buf = generate_docx_bytes()
            st.download_button(label="Descargar Word", data=docx_buf, file_name=base_name + ".docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.success("Archivos listos para descargar.")
    except Exception as e:
        st.error(f"Error al generar los archivos: {e}")

# Automatically save current session data to SQLite database at the end of execution
autosave_current_data()
